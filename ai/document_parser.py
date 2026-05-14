"""
Document Parser — extracts raw text from files.
Local extraction first, Claude Vision as fallback.

Supported formats:
  PDF (text)    -> pdfplumber (free, instant)
  PDF (scanned) -> Tesseract OCR -> Claude Vision
  Excel .xlsx   -> openpyxl (free, instant)
  CSV           -> csv module (free, instant)
  Image JPG/PNG -> Claude Vision API (chargeable)
  Word .docx    -> python-docx (free, instant)
  Text .txt     -> direct read (free, instant)
"""
import os
import csv
import json
import base64
import urllib.request
from pathlib import Path


class PageResult:
    """Result for a single page."""
    def __init__(self, page_num: int, text: str = "", method: str = "",
                 used_claude: bool = False, error: str = ""):
        self.page_num    = page_num
        self.text        = text
        self.method      = method
        self.used_claude = used_claude
        self.error       = error


class ExtractionResult:
    """Full document extraction result."""
    def __init__(self):
        self.pages:        list[PageResult] = []
        self.total_pages:  int  = 0
        self.local_pages:  int  = 0
        self.claude_pages: int  = 0
        self.file_type:    str  = ""
        self.error:        str  = ""

    @property
    def success(self) -> bool:
        return bool(self.pages) and not self.error

    @property
    def full_text(self) -> str:
        return "\n".join(p.text for p in self.pages if p.text)

    @property
    def local_cost_paise(self) -> int:
        return self.local_pages * 10

    @property
    def claude_cost_paise(self) -> int:
        return self.claude_pages * 500

    @property
    def total_cost_paise(self) -> int:
        return self.local_cost_paise + self.claude_cost_paise

    def cost_summary(self) -> str:
        parts = []
        if self.local_pages:
            parts.append(
                f"{self.local_pages} local pages "
                f"@ Rs.0.10 = Rs.{self.local_cost_paise/100:.2f}"
            )
        if self.claude_pages:
            parts.append(
                f"{self.claude_pages} AI pages "
                f"@ Rs.5.00 = Rs.{self.claude_cost_paise/100:.2f}"
            )
        total = self.total_cost_paise / 100
        parts.append(f"Total: Rs.{total:.2f}")
        return " | ".join(parts)


class DocumentParser:

    SUPPORTED = {
        ".pdf":  "pdf",
        ".xlsx": "excel",
        ".xls":  "excel",
        ".csv":  "csv",
        ".txt":  "text",
        ".jpg":  "image",
        ".jpeg": "image",
        ".png":  "image",
        ".docx": "word",
    }

    def __init__(self, api_key: str = "", feature: str = "document_reader"):
        # api_key is now legacy — callers should configure routing in
        # Settings → AI Routing instead. We keep the param so existing call
        # sites don't break; if passed, it's used as a one-shot BYOK
        # override (handy for env-var test runs and the bank-reco fallback
        # path which still surfaces an api-key field today).
        self.api_key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
        self.feature = feature

    @property
    def ai_available(self) -> bool:
        """True if any route (own-key or pooled) can be used right now."""
        if self.api_key:
            return True
        try:
            from core.ai_routing import routing
            from core.license_manager import LicenseManager
            if routing.has_own_key():
                return True
            # Pooled requires a paid license (DEMO/FREE-DEMO bounces out
            # in ai_client.call_messages).
            return LicenseManager().license_key not in (
                "DEMO", "FREE-DEMO", "", None,
            )
        except Exception:
            return False

    def parse(self, filepath: str, sheet_index: int = 0) -> ExtractionResult:
        result = ExtractionResult()
        path   = Path(filepath)

        if not path.exists():
            result.error = f"File not found: {filepath}"
            return result

        ext   = path.suffix.lower()
        ftype = self.SUPPORTED.get(ext)

        if not ftype:
            result.error = (
                f"Unsupported file type: {ext}\n"
                f"Supported: PDF, Excel, CSV, JPG, PNG, Word, TXT"
            )
            return result

        result.file_type = ftype
        handlers = {
            "pdf":   self._parse_pdf,
            "excel": self._parse_excel,
            "csv":   self._parse_csv,
            "text":  self._parse_text,
            "image": self._parse_image,
            "word":  self._parse_word,
        }
        handlers[ftype](filepath, result, sheet_index)
        return result

    # ── PDF ──────────────────────────────────────────────────────────────────

    def _parse_pdf(self, filepath: str, result: ExtractionResult,
                   sheet_index: int = 0):
        try:
            import pdfplumber
        except ImportError:
            result.error = "pdfplumber not installed. Run: pip install pdfplumber"
            return

        try:
            with pdfplumber.open(filepath) as pdf:
                result.total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages[:100], 1):
                    text = page.extract_text() or ""

                    tables = page.extract_tables()
                    if tables:
                        for tbl in tables:
                            for row in tbl:
                                if row:
                                    text += "\n" + " | ".join(
                                        str(c or "") for c in row
                                    )

                    if len(text.strip()) > 50:
                        result.pages.append(PageResult(
                            page_num=i, text=text,
                            method="pdfplumber", used_claude=False
                        ))
                        result.local_pages += 1
                    else:
                        page_result = self._ocr_page(page, i)
                        result.pages.append(page_result)
                        if page_result.used_claude:
                            result.claude_pages += 1
                        else:
                            result.local_pages += 1

        except Exception as e:
            result.error = f"PDF error: {e}"

    def _ocr_page(self, pdf_page, page_num: int) -> PageResult:
        """Try Tesseract first, Claude Vision fallback."""
        try:
            import pytesseract
            img  = pdf_page.to_image(resolution=200)
            text = pytesseract.image_to_string(
                img.original, lang="eng", config="--psm 6"
            )
            if len(text.strip()) > 30:
                return PageResult(
                    page_num=page_num, text=text,
                    method="tesseract_ocr", used_claude=False
                )
        except Exception:
            pass

        if self.ai_available:
            try:
                import io
                img = pdf_page.to_image(resolution=150)
                buf = io.BytesIO()
                img.original.save(buf, format="JPEG")
                b64  = base64.b64encode(buf.getvalue()).decode()
                text = self._claude_vision(b64, "image/jpeg")
                return PageResult(
                    page_num=page_num, text=text,
                    method="claude_vision", used_claude=True
                )
            except Exception as e:
                return PageResult(page_num=page_num, error=str(e), used_claude=True)

        return PageResult(
            page_num=page_num,
            error=(
                "Scanned page — Tesseract not installed and AI routing not "
                "configured. Open Settings → AI Routing to enable AI parsing."
            ),
        )

    # ── Excel ─────────────────────────────────────────────────────────────────

    def _parse_excel(self, filepath: str, result: ExtractionResult,
                     sheet_index: int = 0):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, data_only=True)
            result.total_pages = len(wb.worksheets)
            ws = wb.worksheets[min(sheet_index, len(wb.worksheets) - 1)]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if any(c is not None for c in row):
                    rows.append(" | ".join(
                        str(c) if c is not None else "" for c in row
                    ))
            result.pages.append(PageResult(
                page_num=1, text="\n".join(rows),
                method="openpyxl", used_claude=False
            ))
            result.local_pages += 1
        except Exception as e:
            result.error = f"Excel error: {e}"

    # ── CSV ───────────────────────────────────────────────────────────────────

    def _parse_csv(self, filepath: str, result: ExtractionResult,
                   sheet_index: int = 0):
        try:
            rows = []
            for enc in ["utf-8", "latin-1", "cp1252"]:
                try:
                    with open(filepath, encoding=enc, newline="") as f:
                        rows = list(csv.reader(f))
                    break
                except UnicodeDecodeError:
                    continue
            result.total_pages = 1
            result.pages.append(PageResult(
                page_num=1,
                text="\n".join(" | ".join(row) for row in rows),
                method="csv", used_claude=False
            ))
            result.local_pages += 1
        except Exception as e:
            result.error = f"CSV error: {e}"

    # ── Text ──────────────────────────────────────────────────────────────────

    def _parse_text(self, filepath: str, result: ExtractionResult,
                    sheet_index: int = 0):
        try:
            text = ""
            for enc in ["utf-8", "latin-1", "cp1252"]:
                try:
                    with open(filepath, encoding=enc) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            result.total_pages = 1
            result.pages.append(PageResult(
                page_num=1, text=text, method="text", used_claude=False
            ))
            result.local_pages += 1
        except Exception as e:
            result.error = f"Text error: {e}"

    # ── Word ──────────────────────────────────────────────────────────────────

    def _parse_word(self, filepath: str, result: ExtractionResult,
                    sheet_index: int = 0):
        try:
            import docx
            doc   = docx.Document(filepath)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            result.total_pages = 1
            result.pages.append(PageResult(
                page_num=1, text="\n".join(parts),
                method="python-docx", used_claude=False
            ))
            result.local_pages += 1
        except Exception as e:
            result.error = f"Word error: {e}"

    # ── Image ─────────────────────────────────────────────────────────────────

    def _parse_image(self, filepath: str, result: ExtractionResult,
                     sheet_index: int = 0):
        if not self.ai_available:
            result.error = (
                "AI routing not configured.\n"
                "Open Settings → AI Routing to paste your Anthropic key or "
                "enable pooled credits."
            )
            return
        try:
            from PIL import Image
            import io
            img = Image.open(filepath)
            if max(img.size) > 2000:
                img.thumbnail((2000, 2000))
            buf  = io.BytesIO()
            fmt  = "JPEG" if filepath.lower().endswith((".jpg", ".jpeg")) else "PNG"
            img.save(buf, format=fmt)
            b64  = base64.b64encode(buf.getvalue()).decode()
            mime = "image/jpeg" if fmt == "JPEG" else "image/png"
            text = self._claude_vision(b64, mime)
            result.total_pages = 1
            result.pages.append(PageResult(
                page_num=1, text=text,
                method="claude_vision", used_claude=True
            ))
            result.claude_pages += 1
        except Exception as e:
            result.error = f"Image error: {e}"

    # ── Claude Vision ─────────────────────────────────────────────────────────

    def _claude_vision(self, b64: str, mime: str) -> str:
        payload = {
            "model": "claude-sonnet-4-6",
            "max_tokens": 4096,
            "messages": [{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": mime,
                            "data": b64,
                        },
                    },
                    {
                        "type": "text",
                        "text": (
                            "Extract ALL text and transaction data from this document. "
                            "For each transaction include: date, description, debit amount, "
                            "credit amount, balance. Preserve table structure."
                        ),
                    },
                ],
            }],
        }
        # Legacy api_key override: when a caller still passes an explicit
        # api_key (e.g. the bank-reco fallback before Phase 2b lands), we
        # hit Anthropic directly so the legacy code path isn't disturbed.
        if self.api_key:
            req = urllib.request.Request(
                "https://api.anthropic.com/v1/messages",
                data=json.dumps(payload).encode(),
                headers={
                    "Content-Type":      "application/json",
                    "x-api-key":         self.api_key,
                    "anthropic-version": "2023-06-01",
                },
            )
            with urllib.request.urlopen(req, timeout=60) as resp:
                data = json.loads(resp.read())
            return data["content"][0]["text"]

        # Routed call — honours user's Settings → AI Routing choice.
        from ai.ai_client import call_messages
        data = call_messages(self.feature, payload, timeout=60)
        return data["content"][0]["text"]
