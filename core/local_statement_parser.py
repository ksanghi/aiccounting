"""
Local Document Parser — heuristic, format-agnostic financial-document reader.

Reads CSV/Excel/PDF/TXT/DOCX with stdlib + pdfplumber + openpyxl + python-docx.
No paid AI calls.

Designed for re-use across:

  • Bank reconciliation   →  parse_bank_statement(path)
  • Ledger reconciliation →  parse_ledger(path) [later]
  • Trial balance import  →  parse_trial_balance(path) [later]

The shared primitives are file-reading (_read_csv / _read_excel / _read_pdf
/ _read_docx / _read_text) and table-finding (_lines_from_table with a
configurable HeaderSchema). Per-consumer logic lives in the parse_* methods.

Image (.jpg/.jpeg/.png) and scanned PDFs without extractable text return
success=False with an explanatory error — the caller is expected to ask the
user before falling back to Claude (paid).
"""
from __future__ import annotations

import csv
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional


# ── Known Indian banks — order matters: longer names first ──────────────────

_BANK_NAMES = [
    "State Bank of India", "Punjab National Bank", "Bank of Baroda",
    "Bank of India", "Bank of Maharashtra", "Central Bank of India",
    "Kotak Mahindra Bank", "Standard Chartered", "Deutsche Bank",
    "South Indian Bank", "Karnataka Bank", "Karur Vysya Bank",
    "Tamilnad Mercantile Bank", "Federal Bank", "Union Bank of India",
    "Indian Bank", "Indian Overseas Bank", "UCO Bank",
    "IDFC FIRST Bank", "IDFC FIRST", "IDFC", "IDBI Bank", "IDBI",
    "IndusInd Bank", "IndusInd", "Yes Bank", "RBL Bank", "RBL",
    "DBS Bank", "DBS", "HSBC", "Citibank", "Citi",
    "DCB Bank", "DCB",
    "HDFC Bank", "HDFC", "ICICI Bank", "ICICI", "Axis Bank", "Axis",
    "Kotak", "PNB", "SBI", "BOI", "BOB",
]

# Account number: label + 9-18 digit cluster (incl. masked XXXX1234).
_ACCOUNT_LABEL_RE = re.compile(
    r"""(?ix)
    \b
    (?: a/?c \s* (?:no\.?|number|num|\#)?
      | account \s* (?:no\.?|number|num|\#)
      | acct \s* (?:no\.?|number|num|\#)
    )
    \s* [:\-]? \s*
    ( [Xx*]{0,12} \d{4,18} )
    """,
)

# Opening / Closing balance: "Opening Balance: 150000.00", "Closing Bal Rs. 1,50,000",
# "Balance Brought Forward: ...", "Balance Carried Forward: ...", with optional
# trailing Cr/Dr indicator.
_OPENING_RE = re.compile(
    r"""(?ix)
    (?:
        opening \s* (?:balance|bal\.?)
      | balance \s* (?:brought \s* forward|b/?f)
      | b/?f \s* balance
    )
    \s* [:\-]? \s*
    (?: rs\.? | inr | ₹ )? \s*
    ( \d[\d,]*(?:\.\d+)? )
    \s* (cr|dr)?
    """,
)
_CLOSING_RE = re.compile(
    r"""(?ix)
    (?:
        closing \s* (?:balance|bal\.?)
      | balance \s* (?:carried \s* forward|c/?f)
      | c/?f \s* balance
    )
    \s* [:\-]? \s*
    (?: rs\.? | inr | ₹ )? \s*
    ( \d[\d,]*(?:\.\d+)? )
    \s* (cr|dr)?
    """,
)

# Period: "From X to Y", "Period: X to Y", "Statement Period X – Y", etc.
_PERIOD_RE = re.compile(
    r"""(?ix)
    (?: from | period \s* [:\-]? | between )
    \s*
    (\d{1,2}[/\-\s\.][A-Za-z0-9]{1,9}[/\-\s\.]\d{2,4} | \d{4}-\d{2}-\d{2})
    \s* (?: \s+ to \s+ | \s* (?: -- | – | — ) \s* | \s+ through \s+ | \s+ till \s+ )
    \s*
    (\d{1,2}[/\-\s\.][A-Za-z0-9]{1,9}[/\-\s\.]\d{2,4} | \d{4}-\d{2}-\d{2})
    """,
)

_DATE_FORMATS = [
    "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d-%b-%Y",
    "%d/%m/%y", "%d-%m-%y", "%Y/%m/%d", "%d %b %Y",
    "%d-%B-%Y", "%d %B %Y", "%m/%d/%Y", "%d.%m.%Y",
    "%d %m %Y", "%d %b %y",
]


def _parse_date(raw: str) -> Optional[str]:
    s = (raw or "").strip()
    if not s:
        return None
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(s, fmt).date().isoformat()
        except ValueError:
            continue
    return None


def _parse_amount(raw) -> Optional[float]:
    if raw is None:
        return None
    s = str(raw).strip()
    if not s:
        return None
    s = re.sub(r"[^\d.\-+]", "", s.replace(",", ""))
    if not s or s in ("-", "+", "."):
        return None
    try:
        return float(s)
    except ValueError:
        return None


# ── Header schema — describes what columns a parser is looking for ──────────

@dataclass
class HeaderSchema:
    """
    Per-consumer column-header dictionary. Each tuple lists candidate header
    strings (case-insensitive, substring match). Empty tuple = column not
    used. Future ledger / trial-balance consumers populate the relevant
    entries (e.g. trial-balance uses `ledger`, `opening`, `closing`).
    """
    date:         tuple[str, ...] = ()
    debit:        tuple[str, ...] = ()
    credit:       tuple[str, ...] = ()
    amount:       tuple[str, ...] = ()
    narration:    tuple[str, ...] = ()
    reference:    tuple[str, ...] = ()
    ledger:       tuple[str, ...] = ()    # trial balance / ledger reco
    opening:      tuple[str, ...] = ()    # trial balance
    closing:      tuple[str, ...] = ()    # trial balance
    voucher_no:   tuple[str, ...] = ()    # ledger reco
    voucher_type: tuple[str, ...] = ()    # ledger reco


# Bank-statement schema (used by parse_bank_statement)
_BANK_SCHEMA = HeaderSchema(
    date=(
        "txn date", "transaction date", "value date", "posting date",
        "date", "trans date", "tran date",
    ),
    debit=(
        "debit", "withdrawal", "withdrawals", "withdrawal amt",
        "withdrawal amount", "dr", "dr amount", "debit amount", "chq amt",
    ),
    credit=(
        "credit", "deposit", "deposits", "deposit amt", "deposit amount",
        "cr", "cr amount", "credit amount",
    ),
    amount=("amount", "txn amount", "transaction amount"),
    narration=(
        "narration", "description", "particulars", "remarks",
        "transaction details", "details", "transaction remarks",
    ),
    reference=(
        "ref no", "ref no.", "reference", "ref", "cheque no", "chq no",
        "chq.no", "chqno", "chequeno", "cheque number",
        "instrument no", "utr", "utr no", "transaction id",
        "txn id", "ref/cheque", "ref/cheque no",
    ),
)


def _pick_header(headers, candidates: tuple[str, ...]) -> Optional[int]:
    if not candidates:
        return None
    norm = [str(h).strip().lower() for h in headers]
    for c in candidates:
        if c in norm:
            return norm.index(c)
    for c in candidates:
        for i, h in enumerate(norm):
            if c in h:
                return i
    return None


# ── Result dataclass ────────────────────────────────────────────────────────

@dataclass
class ParseResult:
    success: bool
    bank_name: Optional[str]         = None
    account_number: Optional[str]    = None
    period_from: Optional[str]       = None       # ISO yyyy-mm-dd
    period_to: Optional[str]         = None
    statement_opening: Optional[float] = None
    statement_closing: Optional[float] = None
    lines: list[dict]                = field(default_factory=list)
    error: Optional[str]             = None
    file_text: str                   = ""         # for AI fallback


# ── Main parser ─────────────────────────────────────────────────────────────

class LocalDocumentParser:
    """
    Format-agnostic financial-document parser. No AI calls.

    Public entry points (one per consumer):
        parse_bank_statement(path)  → ParseResult with bank_name, account
                                       number, period, transaction lines.
        parse_ledger(path)          → planned
        parse_trial_balance(path)   → planned

    Internal primitives (reusable across consumers):
        _read_any(path)             → (tables, raw_text)
        _detect_bank_name / account / period
        _lines_from_table(tbl, schema, sign_resolver)
    """

    # ── Public consumers ──────────────────────────────────────────────────

    def parse_bank_statement(self, file_path: str) -> ParseResult:
        """Bank-statement consumer: finds transaction table with debit/credit/amount."""
        meta = self._extract_with_meta(file_path)
        if meta.get("error_result"):
            return meta["error_result"]
        tables, text = meta["tables"], meta["text"]

        bank_name      = self._detect_bank_name(text)
        account_number = self._detect_account_number(text)
        period_from, period_to = self._detect_period(text)
        opening = self._detect_balance(text, _OPENING_RE)
        closing = self._detect_balance(text, _CLOSING_RE)

        lines = self._extract_lines(
            tables,
            schema=_BANK_SCHEMA,
            row_to_line=self._bank_row_to_line,
        )
        if not lines:
            return ParseResult(
                success=False,
                bank_name=bank_name,
                account_number=account_number,
                period_from=period_from,
                period_to=period_to,
                statement_opening=opening,
                statement_closing=closing,
                file_text=text,
                error=(
                    "Could not locate a transaction table. The file may be a "
                    "scanned image, or use a layout the local parser doesn't "
                    "recognise. Try the AI parser."
                ),
            )

        return ParseResult(
            success=True,
            bank_name=bank_name,
            account_number=account_number,
            period_from=period_from,
            period_to=period_to,
            statement_opening=opening,
            statement_closing=closing,
            lines=lines,
            file_text=text,
        )

    # ── Format dispatch (shared primitive) ───────────────────────────────

    def _extract_with_meta(self, file_path: str) -> dict:
        """
        Read the file into (tables, text). On failure, returns a dict with
        an `error_result` ParseResult. Used by every parse_* consumer.
        """
        path = Path(file_path)
        if not path.exists():
            return {"error_result": ParseResult(
                success=False, error=f"File not found: {file_path}"
            )}

        ext = path.suffix.lower()
        try:
            if ext == ".csv":
                tables, text = self._read_csv(path)
            elif ext in (".xlsx", ".xls"):
                tables, text = self._read_excel(path)
            elif ext == ".pdf":
                tables, text = self._read_pdf(path)
            elif ext == ".txt":
                tables, text = [], path.read_text(encoding="utf-8", errors="replace")
            elif ext == ".docx":
                tables, text = self._read_docx(path)
            elif ext in (".jpg", ".jpeg", ".png"):
                return {"error_result": ParseResult(
                    success=False,
                    error=(
                        "Image files cannot be parsed locally. "
                        "Use the AI parser for scanned statements."
                    ),
                )}
            else:
                return {"error_result": ParseResult(
                    success=False,
                    error=f"Unsupported file type: {ext}",
                )}
        except Exception as e:
            return {"error_result": ParseResult(
                success=False, error=f"Could not read file: {e}"
            )}

        return {"tables": tables, "text": text}

    # ── Per-format readers ──────────────────────────────────────────────────

    def _read_csv(self, path: Path):
        for enc in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
            try:
                with open(path, newline="", encoding=enc) as f:
                    rows = [list(r) for r in csv.reader(f)]
                text = "\n".join(",".join(str(c) for c in r) for r in rows)
                return [rows], text
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode CSV with utf-8/latin-1/cp1252.")

    def _read_excel(self, path: Path):
        # openpyxl handles .xlsx but explicitly rejects the legacy .xls
        # binary format — many banks still email .xls. Route .xls through
        # xlrd 1.2.0 (the last release that supports it).
        if path.suffix.lower() == ".xls":
            return self._read_xls_legacy(path)

        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True, read_only=True)
        tables = []
        text_parts: list[str] = []
        for sheet in wb.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                rows.append(["" if c is None else str(c) for c in row])
            tables.append(rows)
            text_parts.append("\n".join(",".join(r) for r in rows))
        return tables, "\n\n".join(text_parts)

    def _read_xls_legacy(self, path: Path):
        try:
            import xlrd
        except ImportError as e:
            raise RuntimeError(
                "This is a legacy .xls file. Install xlrd to read it "
                "(pip install \"xlrd==1.2.0\"), or re-save the file as "
                ".xlsx / .csv."
            ) from e

        book = xlrd.open_workbook(str(path))
        tables: list[list[list[str]]] = []
        text_parts: list[str] = []
        for sheet in book.sheets():
            rows: list[list[str]] = []
            for r in range(sheet.nrows):
                cells = sheet.row_values(r)
                rows.append(["" if c is None else str(c) for c in cells])
            tables.append(rows)
            text_parts.append("\n".join(",".join(r) for r in rows))
        return tables, "\n\n".join(text_parts)

    def _read_pdf(self, path: Path):
        import pdfplumber
        tables = []
        text_parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
                for tbl in page.extract_tables() or []:
                    norm = [["" if c is None else str(c) for c in row] for row in tbl]
                    if norm:
                        tables.append(norm)
        return tables, "\n\n".join(text_parts)

    def _read_docx(self, path: Path):
        from docx import Document
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        tables = []
        for tbl in doc.tables:
            rows = []
            for row in tbl.rows:
                rows.append([c.text for c in row.cells])
            tables.append(rows)
        return tables, text

    # ── Meta detection (reusable; also called by the AI fallback path) ──────

    def extract_meta_from_text(self, text: str) -> dict:
        """
        Run the meta-detection regexes on raw document text and return
        whatever is found. Public so the AI fallback in BankReconciler can
        validate statement ownership using the same heuristics as the
        local parser, without duplicating the regex.
        """
        period_from, period_to = self._detect_period(text)
        return {
            "bank_name":         self._detect_bank_name(text),
            "account_number":    self._detect_account_number(text),
            "period_from":       period_from,
            "period_to":         period_to,
            "statement_opening": self._detect_balance(text, _OPENING_RE),
            "statement_closing": self._detect_balance(text, _CLOSING_RE),
        }

    def _detect_bank_name(self, text: str) -> Optional[str]:
        lower = text.lower()
        for bank in _BANK_NAMES:
            if bank.lower() in lower:
                return bank
        return None

    def _detect_account_number(self, text: str) -> Optional[str]:
        m = _ACCOUNT_LABEL_RE.search(text)
        if m:
            return m.group(1).strip()
        return None

    def _detect_period(self, text: str) -> tuple[Optional[str], Optional[str]]:
        m = _PERIOD_RE.search(text)
        if not m:
            return None, None
        return _parse_date(m.group(1)), _parse_date(m.group(2))

    @staticmethod
    def _detect_balance(text: str, regex: re.Pattern) -> Optional[float]:
        """
        Extract opening/closing balance. Treats trailing 'Cr' as positive
        (bank's natural balance), 'Dr' as negative (overdraft).
        """
        m = regex.search(text)
        if not m:
            return None
        raw = m.group(1).replace(",", "")
        try:
            val = float(raw)
        except ValueError:
            return None
        suffix = (m.group(2) or "").lower() if m.lastindex and m.lastindex >= 2 else ""
        if suffix == "dr":
            val = -val
        return val

    # ── Generic line extraction ─────────────────────────────────────────────

    def _extract_lines(
        self,
        tables,
        schema: HeaderSchema,
        row_to_line,
    ) -> list[dict]:
        """
        Iterate tables, find the first that looks like a transaction layout
        (matches the schema), and yield lines via the row_to_line callable.
        """
        for tbl in tables:
            lines = self._lines_from_table(tbl, schema, row_to_line)
            if lines:
                return lines
        return []

    def _lines_from_table(
        self,
        rows: list[list[str]],
        schema: HeaderSchema,
        row_to_line,
    ) -> list[dict]:
        if not rows:
            return []

        # A row is a header if it has the date column AND at least one of the
        # numeric columns the schema cares about (debit / credit / amount /
        # opening / closing). Search the first 30 rows.
        numeric_groups = (
            schema.debit, schema.credit, schema.amount,
            schema.opening, schema.closing,
        )
        header_idx = None
        for i, row in enumerate(rows[:30]):
            if _pick_header(row, schema.date) is None:
                continue
            if any(
                _pick_header(row, g) is not None for g in numeric_groups
            ):
                header_idx = i
                break
        if header_idx is None:
            return []

        headers = rows[header_idx]
        body    = rows[header_idx + 1:]

        col_idx = {
            "date":         _pick_header(headers, schema.date),
            "debit":        _pick_header(headers, schema.debit),
            "credit":       _pick_header(headers, schema.credit),
            "amount":       _pick_header(headers, schema.amount),
            "narration":    _pick_header(headers, schema.narration),
            "reference":    _pick_header(headers, schema.reference),
            "ledger":       _pick_header(headers, schema.ledger),
            "opening":      _pick_header(headers, schema.opening),
            "closing":      _pick_header(headers, schema.closing),
            "voucher_no":   _pick_header(headers, schema.voucher_no),
            "voucher_type": _pick_header(headers, schema.voucher_type),
        }

        lines: list[dict] = []
        for line_index, row in enumerate(body):
            if not any(str(c).strip() for c in row):
                continue
            ln = row_to_line(row, col_idx, line_index)
            if ln is not None:
                lines.append(ln)
        return lines

    # ── Bank-statement row mapper ───────────────────────────────────────────

    def _bank_row_to_line(
        self,
        row: list[str],
        col_idx: dict,
        line_index: int,
    ) -> Optional[dict]:
        idx_date = col_idx["date"]
        if idx_date is None or idx_date >= len(row):
            return None
        txn_date = _parse_date(str(row[idx_date]))
        if not txn_date:
            return None

        sign, amount = self._row_sign_amount(
            row, col_idx["debit"], col_idx["credit"], col_idx["amount"],
        )
        if sign is None or amount is None or amount <= 0:
            return None

        narration = (
            str(row[col_idx["narration"]]).strip()
            if col_idx["narration"] is not None and col_idx["narration"] < len(row)
            else ""
        )
        reference = (
            str(row[col_idx["reference"]]).strip()
            if col_idx["reference"] is not None and col_idx["reference"] < len(row)
            else ""
        )
        return {
            "line_index": line_index,
            "txn_date":   txn_date,
            "amount":     round(amount, 2),
            "sign":       sign,
            "narration":  narration,
            "reference":  reference,
            "raw_row":    list(row),
        }

    @staticmethod
    def _row_sign_amount(
        row: list[str],
        idx_debit: Optional[int],
        idx_credit: Optional[int],
        idx_amount: Optional[int],
    ) -> tuple[Optional[str], Optional[float]]:
        debit = (
            _parse_amount(row[idx_debit])
            if idx_debit is not None and idx_debit < len(row)
            else None
        )
        credit = (
            _parse_amount(row[idx_credit])
            if idx_credit is not None and idx_credit < len(row)
            else None
        )
        if debit and debit > 0:
            return "DR", debit
        if credit and credit > 0:
            return "CR", credit
        if idx_amount is not None and idx_amount < len(row):
            amt = _parse_amount(row[idx_amount])
            if amt is not None and amt != 0:
                if amt < 0:
                    return "DR", abs(amt)
                return "CR", amt
        return None, None
